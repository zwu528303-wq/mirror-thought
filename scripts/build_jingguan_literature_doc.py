from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT_PATH = Path("docs/镜观_文献方法转译确认表.docx")

FONT_EAST_ASIA = "Microsoft YaHei"
FONT_LATIN = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY_FILL = "F2F4F7"
LIGHT_BLUE_FILL = "E8EEF5"
BORDER = "B8C2CC"


def set_run_font(run, size=None, bold=False, color=None):
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def set_paragraph_font(paragraph, size=10.5, bold=False, color=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def set_table_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_width(cell, width_inches):
    width = Inches(width_inches)
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_text(cell, text, size=8.5, bold=False, color=None):
    cell.text = ""
    parts = str(text).split("\n")
    for index, part in enumerate(parts):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(part)
        set_run_font(run, size=size, bold=bold, color=color)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    if level == 1:
        paragraph.paragraph_format.space_before = Pt(16)
        paragraph.paragraph_format.space_after = Pt(8)
        size = 16
        color = BLUE
    else:
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
        size = 13
        color = BLUE
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=True, color=color)
    return paragraph


def add_body(doc, text, size=10.5):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    set_run_font(run, size=size)
    return paragraph


def format_table(table, widths, header_fill=GRAY_FILL, font_size=8.5):
    table.autofit = False
    set_table_borders(table)
    set_table_width(table, int(sum(widths) * 1440))
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        set_row_cant_split(row)
        for cell_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_width(cell, widths[cell_index])
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_paragraph_font(paragraph, size=font_size, bold=True, color=DARK_BLUE)
            else:
                for paragraph in cell.paragraphs:
                    set_paragraph_font(paragraph, size=font_size)


def build_document():
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("镜观：文献方法转译确认表")
    set_run_font(run, size=22, bold=True, color=DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run("用于小组确认：哪些文献原则可以进入 system prompt / skill / eval")
    set_run_font(run, size=11, color=RGBColor(80, 80, 80))

    meta_table = doc.add_table(rows=1, cols=4)
    meta = ["版本", "v0.2 细化确认稿", "日期", "2026-06-06"]
    for i, text in enumerate(meta):
        add_text(meta_table.cell(0, i), text, size=9, bold=i in (0, 2), color=DARK_BLUE if i in (0, 2) else None)
    format_table(meta_table, [0.9, 2.6, 0.9, 2.0], header_fill="FFFFFF", font_size=9)

    add_heading(doc, "使用方式", level=1)
    add_body(
        doc,
        "这份文档不是最终 prompt，也不是直接面向用户的回答规则。它的作用是把前期文献先转译成镜观可讨论、可确认、可实现的产品原则。小组确认后，再把通过的规则写入 system prompt、Anthropic skill 和 eval。",
    )

    flow_table = doc.add_table(rows=2, cols=5)
    flow_headers = ["1 文献阅读", "2 方法提炼", "3 指令草案", "4 小组确认", "5 写入制度"]
    flow_cells = [
        "只提取与镜观相关的方法点。",
        "转成可执行的思想分析方式。",
        "形成 AI 能遵循的规则语言。",
        "判断 OK / 不 OK / 需修改。",
        "进入 prompt、skill、eval 或二期功能。",
    ]
    for i, text in enumerate(flow_headers):
        add_text(flow_table.cell(0, i), text, size=8.5, bold=True, color=DARK_BLUE)
    for i, text in enumerate(flow_cells):
        add_text(flow_table.cell(1, i), text, size=8.2)
    format_table(flow_table, [1.85, 1.95, 1.95, 1.95, 2.1], header_fill=LIGHT_BLUE_FILL, font_size=8.5)

    add_heading(doc, "一、文献方法转译表", level=1)
    add_body(
        doc,
        "这张表是总览。后面的细化规则表会把每篇文献拆成更小的可执行规则，方便逐条确认。",
        size=9.5,
    )

    rows = [
        [
            "Zinaich on Achenbach：超越方法的方法",
            "哲学咨询不能被固定技术流程完全支配，要贴近来访者具体表达。",
            "镜观可以有流程，但不能机械套模板。",
            "第一轮可用结构化框架；第二轮后必须根据用户刚刚说的话选择追问，不能重复模板。",
            "Skill + eval",
            "避免把“超越方法”误解成没有规则；产品仍需要边界和验收标准。",
        ],
        [
            "Lahav 1996：What is Philosophical in Philosophical Counselling?",
            "哲学性在于审查来访者 lived understanding 背后的原则。",
            "镜观不是信息检索，而是帮助用户审查自己已经相信什么。",
            "不回答用户的问题本身；提取其信念、价值、前提和概念关系。",
            "System prompt",
            "用户侧不要出现大段理论解释；避免学术化。",
        ],
        [
            "Raabe 2001：Philosophical Counseling: Theory and Practice",
            "哲学咨询可以有阶段性过程。",
            "镜观对话可分为进入、映射、澄清、张力分析、阶段性结构整理。",
            "根据对话状态标记 phase；达到条件后只生成结构整理，不生成结论。",
            "Schema + prompt",
            "阶段不能太死，否则会像问卷或客服流程。",
        ],
        [
            "2025 LLM 辅助哲学咨询论文",
            "LLM 可用 prompt、fine-tune、RAG，但有信任、隐私、不能真正理解或共情的问题。",
            "MVP 先用 prompt + skill + eval；不假装共情；不默认长期保存完整对话。",
            "不说“我完全理解你”；不建立长期人格档案；保存结构摘要优先于保存原始聊天。",
            "Prompt + 后端隐私规则",
            "后续若做历史记录，必须重新设计隐私同意。",
        ],
        [
            "Chang 2023：Prompting LLMs With the Socratic Method",
            "苏格拉底追问有类型，不是简单问“你觉得呢”。",
            "建立追问动作库：定义、理由、前提、张力、假设排除。",
            "每轮只选择一种追问动作；问题必须指向具体概念、理由、前提或张力。",
            "Skill reference + eval",
            "避免变成机械审问；追问要保持自然语言。",
        ],
        [
            "Socrates 2.0，JMIR 2024",
            "多智能体：回答者、督导、评分员，可降低循环和有害回应。",
            "镜观可做“回答生成 + 规则检查 + 必要时重写”。",
            "若回答含建议、诊断、安慰、多个追问，则判为失败并重写。",
            "Eval + 后端 guardrail",
            "它是 CBT 项目，不要照搬“不合理信念治疗”框架。",
        ],
        [
            "Lahav 1995：Worldview Interpretation",
            "多个具体困惑背后可能共享世界观结构。",
            "当前会话内可观察反复出现的价值标准，如成功、自由、责任、爱。",
            "只在当前会话中暂定记录反复出现的价值标准；MVP 不做跨会话画像。",
            "Skill + 二期功能",
            "跨会话保存很敏感，不能现在就默认做。",
        ],
        [
            "Romizi & Ramharter 2015：Contradictions",
            "哲学咨询要处理矛盾，但不一定立刻消除矛盾。",
            "镜观用“张力”而不是直接说“你矛盾了”。",
            "使用“似乎有一个张力”；先问两边各自依赖什么理由，不急着解决冲突。",
            "Prompt + summary 规则",
            "不能把所有困惑都硬翻译成形式逻辑矛盾。",
        ],
        [
            "Pan / 思想分析相关线索",
            "支持中文语境下“思想分析”这个定位。",
            "镜观使用信念、前提、概念、张力等语言，而不是治疗、症状、疗愈语言。",
            "你是思想分析引擎，不是心理咨询师、人生导师或百科知识库。",
            "Prompt",
            "目前文献清单里只是间接提到，后续最好补 Pan 原文。",
        ],
    ]

    headers = ["文献", "可参考点", "转化成镜观的方法", "AI 指令草案", "落地位置", "风险 / 待确认"]
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        add_text(table.cell(0, i), header, size=8.2, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            add_text(cells[i], text, size=7.7)
    format_table(table, [1.2, 1.65, 1.65, 2.1, 0.9, 2.3], header_fill=GRAY_FILL, font_size=8.2)

    second_section = doc.add_section(WD_SECTION.NEW_PAGE)
    second_section.orientation = WD_ORIENT.LANDSCAPE
    second_section.page_width = Inches(11)
    second_section.page_height = Inches(8.5)
    second_section.top_margin = Inches(0.55)
    second_section.bottom_margin = Inches(0.55)
    second_section.left_margin = Inches(0.55)
    second_section.right_margin = Inches(0.55)
    second_section.header_distance = Inches(0.35)
    second_section.footer_distance = Inches(0.35)

    add_heading(doc, "二、细化规则表：从文献到 AI 微指令", level=1)
    add_body(
        doc,
        "这一部分用于逐条确认。每一条都可以被标记为：采用、改写后采用、暂不采用。",
        size=9.5,
    )

    detailed_headers = ["文献方向", "具体方法点", "镜观可采用的微规则", "可写入 AI 的指令", "禁止误用 / 检查点"]
    detailed_rows = [
        [
            "Achenbach / 超越方法的方法",
            "方法不能压过来访者的具体表达。",
            "第一轮可以有固定骨架；第二轮以后必须贴着用户刚刚说出的词、情境或修正。",
            "不要重复通用模板。每轮追问必须引用或回应用户上一轮的一个具体表达。",
            "若连续两轮出现类似“最需要澄清哪个概念”，判为模板化。",
        ],
        [
            "Achenbach / 超越方法的方法",
            "咨询者不能急于把对方塞进既有分类。",
            "问题域分类只能是暂定辅助，不应先让用户做 A/B/C 选择。",
            "可以说“我暂时把它理解为偏……，但里面也有……，我理解得对吗？”",
            "避免把开场做成客服问卷。",
        ],
        [
            "Achenbach / 超越方法的方法",
            "方法要服务于对话中的发现。",
            "AI 可以承认信息不足，并请求用户补充具体情境。",
            "如果无法形成“惑”，不要硬提取信念；先问具体事件、选择或冲突。",
            "不要从“我很烦”编造家庭、人格、创伤等解释。",
        ],
        [
            "Lahav 1996 / lived understanding",
            "哲学性在于审查来访者活生生的理解。",
            "从用户语言中提取其已经持有的信念、价值、判断、期待和生活标准。",
            "优先问：这个判断在你这里依赖什么标准？这个词在你的处境里是什么意思？",
            "不要把回答变成信息检索或哲学知识讲解。",
        ],
        [
            "Lahav 1996 / 哲学自我审查",
            "审查的是原则和前提，不是经验事实本身。",
            "AI 要区分事实判断、价值判断、概念判断和关系判断。",
            "当用户混合“现实如此”和“应当如此”时，追问两者的边界。",
            "不要替用户验证现实事实，除非产品另设检索功能。",
        ],
        [
            "Lahav 1996 / 智慧取向",
            "目标不是解决单个问题，而是让用户看清自己的思想结构。",
            "结尾不呈现结论，只呈现阶段性结构整理。",
            "总结只能写核心信念、张力、概念歧义、隐含前提、待澄清问题。",
            "禁止写“所以你应该……”“你的选择是……”。",
        ],
        [
            "Raabe / 阶段过程",
            "哲学咨询可以有阶段，但阶段不是僵硬脚本。",
            "镜观可使用 phase：intake、mapping、clarification、tension_analysis、summary。",
            "根据当前材料选择阶段；不要强制每轮都输出完整四段。",
            "若用户只补充一个词，不要强行做完整总结。",
        ],
        [
            "Raabe / 收束机制",
            "对话需要阶段性成果。",
            "第 5 轮后可判断是否能总结；第 8 轮后可主动提示；第 12 轮后必须建议整理。",
            "使用“阶段性结构整理”，不要说“最终结论”。",
            "总结不是建议清单，也不是治疗报告。",
        ],
        [
            "Raabe / 问题域",
            "进入阶段可以识别问题类型。",
            "问题域可作为内部字段：academic、emotional、mixed、work、family、self、unknown。",
            "第一轮可自然确认：“我暂时理解为……，但里面也有……”。",
            "不要要求用户先选 A/B/C；混合问题应允许存在。",
        ],
        [
            "LLM 辅助哲学咨询 / 技术路径",
            "prompt、fine-tune、RAG 各有用途。",
            "MVP 先用 system prompt 固定边界；skill 放方法库；eval 做失败检查。",
            "不要把长文献全部塞进每轮 prompt；先转成被确认的规则。",
            "避免因技术复杂度拖慢 MVP。",
        ],
        [
            "LLM 辅助哲学咨询 / 信任与隐私",
            "AI 无法真正理解或共情，且涉及敏感对话数据。",
            "AI 不应声称完全理解用户；MVP 不默认保存长期完整对话。",
            "使用“如果我理解准确”而不是“我完全懂你”。",
            "不要做长期人格档案，除非未来有明确同意机制。",
        ],
        [
            "LLM 辅助哲学咨询 / 安全边界",
            "危机场景需要中断普通分析。",
            "出现自伤、伤人、告别、具体时间/手段等风险时停止思想分析。",
            "response_type=crisis；不追问信念结构；引导专业支持或紧急服务。",
            "不能把危机内容继续做哲学分析。",
        ],
        [
            "Chang 2023 / 苏格拉底追问分类",
            "追问有类型，不是泛泛问“你觉得呢”。",
            "每轮只选一种追问动作：定义、理由、前提、关系、事实/价值、假设排除。",
            "在输出内部记录 question_move；用户侧只显示自然语言问题。",
            "避免一轮问多个方向。",
        ],
        [
            "Chang 2023 / 定义追问",
            "概念澄清是苏格拉底追问的重要类型。",
            "当用户使用“重要、合适、成功、自由、真诚”等关键词时，优先定义追问。",
            "你这里说的“X”，更指 A，还是 B？",
            "不要问空泛的“你怎么看 X”。",
        ],
        [
            "Chang 2023 / 假设排除",
            "追问可以帮助排除不同解释。",
            "当用户不确定困惑来源时，用二选一或三选一帮助区分可能性。",
            "你担心的更像是实际资源差距，还是人生叙事不成立？",
            "选项必须来自用户文本，不要塞入外部理论。",
        ],
        [
            "Socrates 2.0 / 督导与评分员",
            "LLM 对话容易循环或产生不当回应。",
            "后端可加入 evaluator：检查建议、安慰、诊断、多问题、循环追问。",
            "若检查失败，要求模型重写，并指出失败项。",
            "评价器不是第二个咨询师，不应引入新解释。",
        ],
        [
            "Socrates 2.0 / 与 CBT 的差异",
            "该项目处理“不合理信念”，更接近 CBT。",
            "镜观不能把用户信念标记为不合理或需要纠正。",
            "使用“这个信念与另一个信念似乎有张力”，而不是“这个信念是不合理的”。",
            "不要照搬治疗目标或认知重构语言。",
        ],
        [
            "Lahav 1995 / 世界观诠释",
            "多个困惑背后可能共享世界观结构。",
            "当前会话内可记录反复出现的价值标准，如好生活、成功、自由、责任、爱。",
            "可以说“这个词似乎在你的几个判断里都很关键”。",
            "MVP 不做跨会话积累；跨会话必须另做隐私设计。",
        ],
        [
            "Lahav 1995 / 重新定向",
            "目标不只是解决问题，而是扩大意义视野。",
            "AI 可以帮助用户看到表面问题背后的核心概念。",
            "这里真正需要分析的，可能不是是否选择 A，而是你如何理解“责任”。",
            "不要把“重新定向”写成行动建议。",
        ],
        [
            "Romizi & Ramharter / 矛盾处理",
            "矛盾不一定要被立即指出或消除。",
            "使用“张力”“难以整合”“似乎同时成立”来替代“你矛盾了”。",
            "这里似乎有一个张力：你一方面……，另一方面……。",
            "不要说“你自相矛盾”。",
        ],
        [
            "Romizi & Ramharter / 保留张力",
            "有时应让矛盾继续存在一段时间。",
            "先问两边各自依赖的理由，而不是急着调和。",
            "如果先不急着解决，这两个判断各自依赖什么理由？",
            "不要急着写“其实二者可以兼得”。",
        ],
        [
            "Romizi & Ramharter / 形式逻辑限制",
            "来访者困惑不总是形式逻辑矛盾。",
            "区分逻辑矛盾、价值冲突、概念混淆、事实/价值错位。",
            "这个冲突更像事实判断不一致，还是价值排序难以整合？",
            "不要把所有困惑都硬翻译成 A 与非 A。",
        ],
        [
            "Pan / 思想分析",
            "中文语境下可用“思想分析”定位区别于心理咨询。",
            "用户侧语言应围绕信念、前提、概念、张力、理由，而非治疗、症状、疗愈。",
            "你是思想分析引擎，不是心理咨询师、人生导师或百科知识库。",
            "后续需要补 Pan 原文，确认术语是否准确。",
        ],
    ]

    detailed_table = doc.add_table(rows=1, cols=len(detailed_headers))
    for i, header in enumerate(detailed_headers):
        add_text(detailed_table.cell(0, i), header, size=8.2, bold=True, color=DARK_BLUE)
    for row in detailed_rows:
        cells = detailed_table.add_row().cells
        for i, text in enumerate(row):
            add_text(cells[i], text, size=7.5)
    format_table(detailed_table, [1.45, 1.7, 2.25, 2.3, 2.1], header_fill=GRAY_FILL, font_size=8.2)

    third_section = doc.add_section(WD_SECTION.NEW_PAGE)
    third_section.orientation = WD_ORIENT.LANDSCAPE
    third_section.page_width = Inches(11)
    third_section.page_height = Inches(8.5)
    third_section.top_margin = Inches(0.55)
    third_section.bottom_margin = Inches(0.55)
    third_section.left_margin = Inches(0.55)
    third_section.right_margin = Inches(0.55)
    third_section.header_distance = Inches(0.35)
    third_section.footer_distance = Inches(0.35)

    add_heading(doc, "三、关键措辞对照：推荐说法与避免说法", level=1)

    wording_headers = ["场景", "避免说法", "推荐说法", "原因"]
    wording_rows = [
        [
            "发现矛盾",
            "你这里自相矛盾。",
            "这里似乎有一个张力：你一方面……，另一方面……。",
            "保留暂定性，不把用户放在被纠错的位置。",
        ],
        [
            "用户问该怎么选",
            "你应该选择……",
            "我不能替你做选择，但可以先整理这两个选择背后的信念结构。",
            "保持思想分析边界，不进入行动建议。",
        ],
        [
            "用户寻求安慰",
            "这很正常，你已经做得很好。",
            "我不能直接给安慰式判断，但可以帮你看清你想被确认的是什么。",
            "避免情绪陪伴替代思想分析。",
        ],
        [
            "概念模糊",
            "你怎么看这个问题？",
            "你这里说的“自由”，更指不受限制，还是能按自己的价值排序生活？",
            "把泛泛追问改成具体概念澄清。",
        ],
        [
            "输入太 vague",
            "你可能是因为压力太大。",
            "现在的信息还不足以形成一个“惑”。你愿意先说一个具体情境吗？",
            "不从情绪词编造原因。",
        ],
        [
            "阶段性结束",
            "结论是你需要……",
            "目前整理出的结构是：核心信念、张力、概念歧义、仍需澄清的问题。",
            "结尾只呈现结构，不呈现答案。",
        ],
    ]
    wording_table = doc.add_table(rows=1, cols=len(wording_headers))
    for i, header in enumerate(wording_headers):
        add_text(wording_table.cell(0, i), header, size=9, bold=True, color=DARK_BLUE)
    for row in wording_rows:
        cells = wording_table.add_row().cells
        for i, text in enumerate(row):
            add_text(cells[i], text, size=8.4)
    format_table(wording_table, [1.4, 2.1, 3.4, 2.9], header_fill=LIGHT_BLUE_FILL, font_size=9)

    add_heading(doc, "四、建议先确认的三条产品制度", level=1)

    policy_headers = ["产品制度", "文献支持", "建议处理"]
    policy_rows = [
        [
            "结尾不呈现结论，只做“阶段性结构整理”。",
            "Lahav、Achenbach、矛盾处理文献。",
            "强烈建议采用。输出核心情境、信念、张力、概念歧义、隐含前提和待澄清问题，不输出用户该怎么做。",
        ],
        [
            "AI 可识别问题域，但不要让用户先填 A/B/C。",
            "Raabe 的阶段模型可支持分类，但文献不支持客服式分类。",
            "修改后采用。AI 内部暂定 domain，如 academic / emotional / mixed / work / family / self；第一轮只用自然语言请求确认。",
        ],
        [
            "文献原则必须先经小组确认，再写进 prompt / skill / eval。",
            "LLM 辅助咨询论文、Socrates 2.0 的评估思路。",
            "强烈建议采用。任何文献规则先产出“依据、原则、实现方式、风险”，确认后再施行。",
        ],
    ]
    policy_table = doc.add_table(rows=1, cols=3)
    for i, header in enumerate(policy_headers):
        add_text(policy_table.cell(0, i), header, size=9, bold=True, color=DARK_BLUE)
    for row in policy_rows:
        cells = policy_table.add_row().cells
        for i, text in enumerate(row):
            add_text(cells[i], text, size=8.6)
    format_table(policy_table, [2.7, 2.3, 4.8], header_fill=LIGHT_BLUE_FILL, font_size=9)

    add_heading(doc, "五、建议确认话术", level=1)
    add_body(
        doc,
        "如果小组认可，下一步不是直接让 AI 执行全部文献，而是把上表中通过的项目分别写入：system prompt 的硬边界、Anthropic skill 的方法库、eval 的失败规则和测试案例。",
    )
    add_body(
        doc,
        "建议确认句：我们同意先采用“文献阅读 → 方法转译 → 小组确认 → 写入 prompt / skill / eval → 案例测试”的流程。未经确认的文献观点不直接进入用户侧回答。",
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_document()
    print(OUT_PATH)
